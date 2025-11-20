"""
Microsoft Word document loader (.docx).
"""

from pathlib import Path
from typing import List
import docx

from .base_loader import BaseDocumentLoader, Document
from src.core.errors import DocumentLoadError
from src.core.logger import setup_logger


logger = setup_logger(__name__)


class DOCXLoader(BaseDocumentLoader):
    """Load Microsoft Word (.docx) documents."""
    
    def load(self) -> List[Document]:
        """
        Load DOCX file and extract text from paragraphs.
        
        Returns:
            List containing single Document object
        """
        try:
            doc = docx.Document(self.file_path)
            
            # Extract all paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            if not paragraphs:
                logger.warning(f"Empty DOCX file: {self.file_path.name}")
                return []
            
            text = '\n\n'.join(paragraphs)
            
            metadata = self._get_base_metadata()
            metadata['paragraph_count'] = len(paragraphs)
            metadata['char_count'] = len(text)
            
            # Extract document properties if available
            try:
                core_properties = doc.core_properties
                if core_properties.author:
                    metadata['author'] = core_properties.author
                if core_properties.title:
                    metadata['title'] = core_properties.title
                if core_properties.subject:
                    metadata['subject'] = core_properties.subject
            except:
                pass  # Properties not always available
            
            document = Document(
                text=text.strip(),
                metadata=metadata
            )
            
            logger.info(f"Loaded DOCX file: {self.file_path.name} ({len(paragraphs)} paragraphs)")
            return [document]
            
        except Exception as e:
            logger.error(f"Failed to load DOCX file: {self.file_path.name}", extra={'error': str(e)})
            raise DocumentLoadError(
                f"DOCX file loading failed: {e}",
                details={'path': str(self.file_path), 'error': str(e)}
            )
